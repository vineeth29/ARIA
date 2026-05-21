import os, json, re

SUPPORTED = {
    ".pdf": "pdf",
    ".docx": "word",
    ".doc": "word",
    ".xlsx": "excel",
    ".xls": "excel",
    ".csv": "csv",
    ".txt": "text",
    ".py": "code",
    ".js": "code",
    ".ts": "code",
    ".html": "code",
    ".css": "code",
    ".json": "json",
    ".xml": "text",
    ".md": "text",
    ".log": "text",
}

def is_readable(path):
    ext = os.path.splitext(path)[1].lower()
    return ext in SUPPORTED

def read_pdf(path):
    try:
        import fitz
        doc = fitz.open(path)
        pages = []
        for i, page in enumerate(doc):
            text = page.get_text().strip()
            if text:
                pages.append(f"[Page {i+1}]\n{text}")
        return "\n\n".join(pages)
    except ImportError:
        try:
            import subprocess
            r = subprocess.run(["pdftotext", path, "-"], capture_output=True, text=True)
            if r.returncode == 0:
                return r.stdout
        except Exception:
            pass
        return None, "Install pymupdf: pip install pymupdf"

def read_word(path):
    try:
        from docx import Document
        doc = Document(path)
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts)
    except ImportError:
        return None

def read_excel(path):
    try:
        import openpyxl
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        parts = []
        for sheet in wb.sheetnames:
            ws = wb[sheet]
            parts.append(f"[Sheet: {sheet}]")
            for row in ws.iter_rows(values_only=True):
                row_vals = [str(c) if c is not None else "" for c in row]
                if any(v.strip() for v in row_vals):
                    parts.append(" | ".join(row_vals))
        return "\n".join(parts)
    except ImportError:
        try:
            import csv
            if path.endswith(".csv"):
                with open(path, newline="", encoding="utf-8", errors="ignore") as f:
                    reader = csv.reader(f)
                    return "\n".join(" | ".join(row) for row in reader)
        except Exception:
            pass
        return None

def read_csv(path):
    try:
        import csv
        rows = []
        with open(path, newline="", encoding="utf-8", errors="ignore") as f:
            reader = csv.reader(f)
            for row in reader:
                rows.append(" | ".join(row))
        return "\n".join(rows)
    except Exception as e:
        return None

def read_text(path):
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception:
        return None

def read_json(path):
    try:
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
        return json.dumps(data, indent=2)
    except Exception:
        return read_text(path)

def read_file(path, max_chars=15000):
    path = os.path.expanduser(path)
    if not os.path.exists(path):
        return None, f"File not found: {path}"
    ext = os.path.splitext(path)[1].lower()
    ftype = SUPPORTED.get(ext, "text")
    content = None
    if ftype == "pdf":
        content = read_pdf(path)
    elif ftype == "word":
        content = read_word(path)
    elif ftype == "excel":
        content = read_excel(path)
    elif ftype == "csv":
        content = read_csv(path)
    elif ftype == "json":
        content = read_json(path)
    else:
        content = read_text(path)
    if content is None:
        return None, f"Could not read {ext} file. Missing library."
    if isinstance(content, tuple):
        return content
    if len(content) > max_chars:
        content = content[:max_chars] + f"\n\n[...truncated — {len(content)} chars total]"
    return content, None

def extract_path_from_text(text):
    patterns = [
        r'"([^"]+\.[a-zA-Z]{2,5})"',
        r"'([^']+\.[a-zA-Z]{2,5})'",
        r'([A-Za-z]:\\[^\s,;]+\.[a-zA-Z]{2,5})',
        r'([~/][^\s,;]+\.[a-zA-Z]{2,5})',
        r'(\b[\w\-. ]+\.(?:pdf|docx?|xlsx?|csv|txt|py|js|ts|json|md|log)\b)',
    ]
    for pat in patterns:
        m = re.search(pat, text, re.IGNORECASE)
        if m:
            path = m.group(1).strip()
            return os.path.expanduser(path)
    return None

def is_file_request(text):
    keywords = [
        "read", "open", "summarise", "summarize", "what does", "what's in",
        "explain", "analyse", "analyze", "extract", "find in", "search in",
        "look at", "check", ".pdf", ".docx", ".xlsx", ".csv", ".txt", ".py",
        "this file", "the file", "my file", "this document", "the document"
    ]
    tl = text.lower()
    return any(k in tl for k in keywords) and extract_path_from_text(text) is not None

def get_file_summary_prompt(content, user_question, filename):
    name = os.path.basename(filename)
    return (
        f"[FILE CONTENT — {name}]\n"
        f"{content}\n"
        f"[END OF FILE]\n\n"
        f"User question about this file: {user_question}\n"
        f"Answer based on the file content above. Be specific and accurate."
    )
